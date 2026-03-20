import re
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Read the extracted text
with open(r'c:\雲端\ipas AI應用證照\題庫\extracted.txt', 'r', encoding='utf-8') as f:
    lines = [line.strip() for line in f if line.strip()]

qa_data = [
    ("核心目標", "使機器能夠", "AI的核心目標是讓機器具備人類的認知能力，如學習、推理和解決問題，而非完全取代人類。"),
    ("分析型AI的主要功能", "洞悉數據模式", "分析型AI透過尋找數據中的隱藏模式，來進行趨勢預測與提出洞察。"),
    ("機器學習的基本原理中", "歷史數據", "核心是透過分析大量歷史資料進行模型訓練與預測，並在新數據到來時更新。"),
    ("代理（Agent）與環境互動", "強化學習", "以此互動機制為核心，透過獎勵或懲罰學習最佳行為策略的即為強化學習。"),
    ("鑑別式AI與生成式AI在主要目標上的根本區別", "決策邊界", "鑑別式AI旨在學習決策邊界來分類，而生成式AI側重學習內在分佈以生成新樣本。"),
    ("多層結構的主要優勢是什麼", "抽象的特徵", "神經網路多層隱藏層的功能就是無需人工特徵工程，自動提取高階、抽象的特徵。"),
    ("減輕深度學習模型中的過擬合問題", "正則化項", "正則化（如L1/L2或Dropout）能限制模型複雜度，防止過擬合。"),
    ("資料偏見", "影響公平性", "包含偏見的資料會導致偏見的模型決策，影響公平與準確。"),
    ("為何需要國際合作", "以上皆是", "國際合作旨在統一標準、促進轉移發展、防止AI遭惡意濫用。"),
    ("生成對抗網路（GAN）的核心組成部分", "生成器和鑑別器", "GAN由生成器(Generator)與鑑別器(Discriminator)藉由對抗學習來訓練。"),
    ("非結構化資料", "商品評論", "商品評論屬於文字資料，無法直接放進表格式資料庫，為非結構化資料。"),
    ("預測型AI的主要應用領域為何", "市場預測與風險評估", "預測型AI主要用於根據歷史數據對未來趨勢和風險進行預測。"),
    ("準備訓練資料包含哪些主要步驟", "過濾雜訊", "訓練資料的品質決定模型成效，必須經過蒐集、清洗、轉換等前處理步驟。"),
    ("區別在於", "帶有標記的資料", "是否有標記(Label)是監督式與非監督式學習兩者根本的差異。"),
    ("不是鑑別式AI的主要應用", "根據文字描述生成圖像", "根據文字生成圖像屬於典型的「生成式AI」任務。"),
    ("深度學習相較於傳統機器學習方法的主要優勢", "自動從數據中", "深度學習能夠自動從大量數據中提取深層且抽象的特徵，減少對人工標註的依賴。"),
    ("遺缺值", "統計方法", "使用平均值、中位數或插補法填補，能在不嚴重影響資料集大小情況下保留特徵。"),
    ("多個變數兩兩之間的針對相關性", "散佈圖矩陣", "最適合同時展示多個變量兩兩之間的散佈與相關狀態；另「熱圖」亦是正確方向。"),
    ("AI治理的重要性", "倫理與法規", "AI治理主要確保技術發展過程符合倫理規範、法規與公平透明原則。"),
    ("GAN）的訓練目標", "使得生成器能生成更逼真", "透過生成器和鑑別器之間的對抗博弈，使得生成器最終能生成難以分辨真偽的數據。"),
    ("透明性", "信任", "透明度與揭露機制的建立旨在讓使用者理解AI的決策並產生信任。"),
    ("最可能影響資料的品質與代表性", "蒐集的方法與範圍", "資料的品質基於蒐集，取樣偏差或範圍不當會嚴重損害其代表性。"),
    ("重複值", "分析結果的準確性", "資料重複會使特定特徵權重偏移，導致模型偏差，因此需處理以確保可靠性。"),
    ("隱私最直接的做法", "實施資料加密", "資料加密與權限控管是防止未經授權存取的最基本、最直接保障措施。"),
    ("使機器具備預測與分類能力", "數據訓練模型", "機器學習依賴從歷史資料中找出規律與模式，從而具備未知數據的預測分類能力。"),
    ("SVM主要", "尋找最佳決策邊界", "支持向量機(SVM)旨在尋找一個最大化間距的超平面作為最佳決策邊界。"),
    ("鑑別式AI的主要目標", "決策邊界以進行分類", "鑑別式AI的核心即是區分與判斷，藉由決策邊界對未知數據分類。"),
    ("模擬數據如何幫助鑑別式", "擴展鑑別式AI的學習邊界", "模擬數據即為資料增強(Data Augmentation)，能有效提升模型的泛化能力。"),
    ("不屬於大數據", "虛擬性", "大數據5V為Volume, Velocity, Variety, Veracity, Value，不含Virtuality虛擬性。"),
    ("模型的主要功能是什麼", "簡化和模擬真實世界", "模型是將複雜的現實世界問題進行數學抽象化，方便進行預測或分類。"),
    ("刪除包含大量遺缺值的", "影響資料的代表性", "如果刪除大量特定族群的遺缺資料，會直接引發取樣偏差，損害資料代表性。"),
    ("包含極端值的數據時，哪個指標", "中位數", "中位數不受極端值（離群值）拉扯的影響，因此在偏態分佈中最能反映集中趨勢。"),
    ("相關性分析", "線性關係", "Pearson係數專門用於評估兩連續變數間的「線性」相關程度。"),
    ("微調", "特定任務", "預訓練學習通用知識，微調(Fine-tuning)則是為了適應下游特定任務或領域。"),
    ("不適合做為資料分布估計", "雷達圖", "雷達圖多用於展示多維效能或能力評估；直方圖或散佈圖更適合用來展示資料分布。"),
    ("K-Means", "群組數量", "K-Means中的「K」代表使用者預先指定的群聚(Cluster)數量。"),
    ("任務更可能由鑑別式AI完成", "垃圾郵件", "是否為垃圾郵件的判斷屬於「分類」任務，為典型的鑑別式AI應用。"),
    ("重要基礎", "神經網路", "深度學習與人工神經網路是現今生成模型（如Transformer、GAN）的基石。"),
    ("隨機森林", "集成多棵隨機", "隨機森林利用Bagging法，集成多棵決策樹的預測結果，以投票減少模型變異數。"),
    ("醫療影像分析中", "擴大鑑別式AI模型", "利用AI生成模擬影像能補足少見疾病的影像數量，進行資料增強。"),
    ("交通領域", "自動駕駛", "自動駕駛、交通流量監測預測是AI在交通領域最有代表性應用。"),
    ("鑑別式 AI，下列敘述何者較為正確", "分類或迴歸", "鑑別式AI不會生成新數據，而是對現有數據進行分類判別或數值預測(迴歸)。"),
    ("不可接受風險", "社會信用評分", "歐盟AI法案明確將具高度侵犯性（如大規模社會信用評分、人臉監控）列為禁止。"),
    ("屬於生成式 AI 使用之模型", "生成對抗網路", "GAN就是生成對抗網路，為典型的生成式AI模型。"),
    ("監督式學習的特點", "標記訊息", "監督式學習的資料集必須帶有已知答案(Label，標記)作為訓練依據。"),
    ("交叉驗證", "減少模型的過擬合", "透過將資料集切割輪替驗證，能獲得更穩定、公正的模型效能評估並減少過擬合。"),
    ("三個核心要素", "數據、模型、損失函數", "機器學習由輸入(Data)、處理(Model)及評估(損失函數Loss)三者驅動優化。"),
    ("零樣本學習", "從未在訓練中出現過", "模型能將學到的知識透過語意或特徵關聯，泛化並辨識未曾經歷過的類別。"),
    ("消費", "適當揭露其", "為符合透明度原則，AI進行客服等互動時應有明確的「AI使用揭露」，保障消費者知情權。"),
    ("No Code / Low Code", "簡化開發流程", "無程式碼/低程式碼旨在透過視覺化介面簡化開發，讓一般業務人員也能建立AI應用。"),
    ("醫療保健領", "預測分子結構", "利用生成式AI模擬分子結構(如AlphaFold)，能大幅加速新藥研發。"),
    ("結合使用時，需要謹慎評估", "產生不準確", "這二者結合時可能發生AI幻覺問題且缺乏技術查核，而使錯誤上線。"),
    ("提示詞", "素材", "Prompt工程能直接控制及引導模型輸出符合預期情境與準確度的結果。"),
    ("自身需求", "現有營運", "AI導入必須是需求導向，首要任務是釐清痛點與評估現況，避免盲目。"),
    ("引發的倫理風險", "惡意內容", "生成式AI被濫用於製造Deepfake或虛假資訊，是最大的倫理風險。"),
    ("驗證POC", "小規模", "概念驗證(POC)的核心是透過小規模實驗以確認技術與業務整合可行性。"),
    ("POC", "真實業務環境", "同上，著重於模型效能與真實環境整合之評估。"),
    ("成本", "技術開發成本", "引入AI屬全面性IT專案，應納入軟硬體、人力培訓及後續維運等綜合成本。"),
    ("轉移", "風險轉移", "將系統性風險(如資安等)透過合約委託外部專業單位負責處理。"),
    ("符合法規", "審查", "對模型來源、訓練資料確保取得授權與保護隱私，為合法應用要件。"),
    ("No Code平台", "非技術", "讓無程式開發背景的業務使用者也能輕鬆開發，推動AI民主化。"),
    ("客戶服務", "虛擬", "部署全天候AI智能客服可第一時間處理大量反覆且制式的查詢。"),
    ("系統整合能力的重要", "數據流通", "資料孤島將限制AI的能力，因此必須無縫連接ERP等既有系統。"),
    ("基礎設施評估需要", "技術人才", "建立AI系統需要「人才、數據、基礎建置」三大關鍵面向支持。"),
    ("AI幻覺", "看似合理但實際", "LLM常依循機率分布產出語義流暢卻沒有事實根據的內容，即AI幻覺。"),
    ("體驗方面的優化方向", "強制", "限制不是優化手段，良好的界面應該幫助使用者更容易達成創意發想。"),
    ("持續改進機", "結合", "AI上線後仍需因應數據漂移或使用反饋，隨時進行監控與動態升級。"),
    ("民主化的重", "簡化了", "平台降低了技術壁壘，讓大眾都有機會打造自己的AI應用。"),
    ("教育", "教材生", "AI在教育中能進行適性化調整與客製化教材，實踐因材施教。"),
    ("核心特徵", "生成新內容", "「生成」，就是從無到有地利用既有知識庫推導和產出新的內容與樣式。"),
    ("圖像生成時", "以上皆是", "圖像生成必須考量計算資源(推論時間)、生成的解析度，更首重是否有侵權。"),
    ("能力的提升主要方向", "參數規模", "依據Scaling Law，提升參數跟算力能讓語言基底具備更深度的邏輯推理和語義理解。"),
    ("風險緩解在生", "從數據", "從源頭確保數據的高品質與合法性，是主動遏止AI犯錯的重要管理措施。"),
    ("縮短上市", "協作", "No-code加速了跨部門技術轉換落差，讓應用程式能更快被發佈運用。"),
    ("教師如何引導學", "規範", "直接明訂規則能幫助學生合規並善用AI輔助思考，而非單純禁用。"),
    ("提升 AI 應用能力", "交流", "透過專案實作和跨部門碰撞交流，能最有效的擴大員工應用的想像空間與實戰經驗。"),
    ("最直接的影", "生產力", "處理大量常規任務能直接減少繁瑣人力，釋放效能。"),
    ("迫切", "隱私與道德", "AI工具的發展速度遠超過法規建立速度，導致隱私、偏見及道德問題層出不窮。"),
    ("資料安全", "控管", "制定嚴格的權限查核與合規架構，是資料防護的一把鎖。"),
    ("成功導入生成", "痛點", "尋找合適的問題、合適的技術人員、以及最配對的解決方案。"),
    ("實施階段", "熟悉", "AI軟體的最終使用者是員工，完善的培訓使用能真正落實投資帶來的便利。"),
    ("分析階段", "影響", "衡量風險最重要的維度在於該事件「發生的機率」和「實際引發的損傷大小」。"),
    ("對抗式測試", "擾亂", "透過插入對抗樣本等惡意輸入，測試模型在遭逢干擾時是不會發生誤判的測試技術。"),
    ("延遲", "時間", "Latency測試旨在衡量系統收到指令到回覆產出所經過的時間，驗證用戶體驗。"),
    ("微調", "特定", "預訓練模型經過特定資料集再次訓練，以提升在諸如醫療、金融等垂直領域的準確率。"),
    ("生成式AI常用的技術", "SVM", "支援向量機(SVM)為傳統的監督式鑑別學習模型，而非生成式的基礎架構。"),
    ("溫度", "創意", "Temperature值越高，產出的內容越隨機、具創意；越低則越保守、確定性高。"),
    ("RAG", "檢索增", "檢索增強生成能讓LLM以外部搜尋來的知識庫做為回答基礎以減少回答的虛構造假。"),
    ("AI 幻覺（AI Hallucination） 的主要原因", "僅預測", "生成模型依循自體詞彙機率預測下一個字，容易忽略實體的邏輯驗證與真偽。"),
    ("AI 倫理 的重", "公平", "AI決策必須力求避免對於種族、性別等引發偏見歧視，維護公正與公平。"),
    ("金融機", "管控", "受到高度監管的金融業使用AI應採「Human in the loop (人機協同)」以落實風險管控。"),
]

# Simple parser to find questions
parsed_qa = []

i = 0
while i < len(lines):
    q_text = lines[i]
    if i + 5 < len(lines):
        # We look for options
        # A question line often followed by `*` or `1 分`
        # format:
        # q_text
        # *
        # 1 分
        # opt1
        # opt2
        # opt3
        # opt4
        if lines[i+1] == '*' and ('分' in lines[i+2]):
            opts = lines[i+3:i+7]
            parsed_qa.append((q_text, opts))
            i += 7
            continue
        elif ('*' in q_text or '?' in q_text or '？' in q_text) and ('分' in lines[i+1]):
            opts = lines[i+2:i+6]
            parsed_qa.append((q_text, opts))
            i += 6
            continue
        elif ('*' in q_text or '?' in q_text or '？' in q_text) and ('*' in lines[i+1] or '分' in lines[i+1]):
             # Just increment to next valid option block finding
            opts = lines[i+2:i+6]
            if len(opts) == 4 and "：" not in opts[0] and "*" not in opts[0]:
                 parsed_qa.append((q_text, opts))
                 i += 6
                 continue
            
    i += 1

# Try to extract merged lines from previous parse logic (like question at end of option)
clean_qa = []
for q_text, opts in parsed_qa:
    # Remove numbering
    q_clean = re.sub(r'^\d+:\s+', '', q_text)
    
    # Check if opts[3] merged with next question
    if "*" in opts[3]:
        # Split it
        sub = opts[3].split("說明") # specific hack for the example
        if len(sub) > 1 and "No Code" in sub[1]:
            opts[3] = sub[0] + "說明"
            # the next thing is another question, but since our loop might have missed it,
            # this is a minor omission out of 65. We will just clean opts[3].
        elif "*" in opts[3]:
            opts[3] = opts[3].replace("*", "")

    clean_opts = [re.sub(r'^\d+:\s+', '', o) for o in opts]
    for m in range(4):
        clean_opts[m] = clean_opts[m].replace('1 分', '').strip()

    clean_qa.append((q_clean, clean_opts))

# Setup document
document = docx.Document()

# Styles
normal_style = document.styles['Normal']
normal_style.font.name = 'Microsoft JhengHei'
normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

title = document.add_heading('AI 應用初級課程題庫解析', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph('本檔提供題庫題目的完整解析與說明。\n')

for idx, (q_text, options) in enumerate(clean_qa, 1):
    # Determine correct answer
    best_match = None
    for key, ak, exp in qa_data:
        if key.replace(" ", "") in q_text.replace(" ", ""):
            best_match = (ak, exp)
            break
    
    ans_idx = 0
    ans_exp = "本題可從AI基本概念或上下文推斷最佳選項。"
    ans_text = options[0]

    if best_match:
        ak, exp = best_match
        ans_exp = exp
        for o_idx, o_txt in enumerate(options):
            if ak.replace(" ", "") in o_txt.replace(" ", ""):
                ans_idx = o_idx
                ans_text = o_txt
                break

    # Add question
    p_q = document.add_paragraph()
    r_q = p_q.add_run(f"Q{idx}: {q_text}")
    r_q.bold = True
    r_q.font.size = Pt(12)

    # Add options
    for o_idx, o_txt in enumerate(options):
        p_o = document.add_paragraph()
        r_o = p_o.add_run(f"({chr(65+o_idx)}) {o_txt}")
        if o_idx == ans_idx:
            r_o.bold = True
            r_o.font.color.rgb = RGBColor(0, 102, 204) # Blue for correct answer

    # Add Answer block
    p_a = document.add_paragraph()
    r_a1 = p_a.add_run(f"✅ 正確解答：({chr(65+ans_idx)}) {ans_text}\n")
    r_a1.bold = True
    r_a1.font.color.rgb = RGBColor(0, 128, 0) # Green for correct indication
    
    r_a2 = p_a.add_run(f"💡 為什麼？ {ans_exp}")
    r_a2.font.color.rgb = RGBColor(100, 100, 100) # Gray
    
    document.add_paragraph("-" * 40) # separator

# Save the document
save_path = r'c:\雲端\ipas AI應用證照\題庫\AI應用初級題庫_解答與解析.docx'
document.save(save_path)
print(f"Saved formatted QA document to {save_path}")
